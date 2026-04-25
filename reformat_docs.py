"""
Reformat Chinese academic documents to meet strict formatting requirements.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SIMSUN = "宋体"
SIMHEI = "黑体"
TIMES_NEW_ROMAN = "Times New Roman"

def set_font(run, name_cn, name_en=None, size_pt=12, bold=False):
    run.font.name = name_cn
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._r
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name_cn)
    rfonts.set(qn("w:ascii"), name_en or name_cn)
    rfonts.set(qn("w:hAnsi"), name_en or name_cn)

def set_para_format(para, line_spacing=1.5, space_before=0, space_after=0,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=False):
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Pt(24)

def add_body_para(doc, text, font_name=SIMSUN, font_size=12, indent_first=True,
                  bold=False, en_font=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    set_para_format(para, indent_first=indent_first, alignment=alignment)
    run = para.add_run(text)
    set_font(run, font_name, en_font or TIMES_NEW_ROMAN, font_size, bold)
    return para

def add_heading1(doc, text):
    para = doc.add_paragraph()
    set_para_format(para, line_spacing=1.5, space_before=6, space_after=6,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False)
    run = para.add_run(text)
    set_font(run, SIMHEI, TIMES_NEW_ROMAN, 15, bold=True)
    # Apply Heading 1 outline level
    pPr = para._p.get_or_add_pPr()
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), "0")
    pPr.append(outlineLvl)
    return para

def add_heading2(doc, text):
    para = doc.add_paragraph()
    set_para_format(para, line_spacing=1.5, space_before=4, space_after=4,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False)
    run = para.add_run(text)
    set_font(run, SIMHEI, TIMES_NEW_ROMAN, 14, bold=True)
    pPr = para._p.get_or_add_pPr()
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), "1")
    pPr.append(outlineLvl)
    return para

def add_heading_unnumbered(doc, text, size=15):
    """Like heading1 but unnumbered (for 摘要, 目录, 参考文献)."""
    para = doc.add_paragraph()
    set_para_format(para, line_spacing=1.5, space_before=6, space_after=6,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    run = para.add_run(text)
    set_font(run, SIMHEI, TIMES_NEW_ROMAN, size, bold=True)
    pPr = para._p.get_or_add_pPr()
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), "0")
    pPr.append(outlineLvl)
    return para

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())
    return para

def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br

def insert_page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pgBr = OxmlElement("w:pageBreakBefore")
    pgBr.set(qn("w:val"), "1")
    pPr.append(pgBr)

def add_toc_field(doc):
    """Add a Word TOC field that updates when opened."""
    para = doc.add_paragraph()
    set_para_format(para, indent_first=False)
    run = para.add_run()
    set_font(run, SIMSUN, TIMES_NEW_ROMAN, 12)
    # Build the TOC field XML
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)
    return para

def set_page_margins(doc, margin_cm=2.5):
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


# ============================================================
# Build 文献综述.docx
# ============================================================
def build_wenxian_zongsu():
    src = Document("/home/runner/work/dpft/dpft/文献综述.docx")
    paras = [p.text.strip() for p in src.paragraphs]

    # Collect full reference texts
    refs = []
    in_refs = False
    for t in paras:
        if t == "参考文献":
            in_refs = True
            continue
        if in_refs and t:
            refs.append(t)

    # Collect chapter content paragraphs (index by chapter)
    ch1_paras = [paras[4], paras[5], paras[6]]
    sec21 = paras[9]
    sec22_paras = [paras[11], paras[12]]
    sec23 = paras[14]
    sec31 = [paras[17], paras[18]]
    sec32 = [paras[20], paras[21], paras[22], paras[23]]
    sec33 = [paras[25], paras[26]]
    sec41 = [paras[29], paras[30], paras[31]]
    sec42 = paras[33]
    ch5_paras = [paras[35], paras[36]]

    doc = Document()
    set_page_margins(doc)

    # ---- Cover ----
    p = doc.add_paragraph()
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r = p.add_run("文献综述")
    set_font(r, SIMHEI, TIMES_NEW_ROMAN, 22, bold=True)  # 小二 = 18pt in some, 22pt for 小二

    p2 = doc.add_paragraph()
    set_para_format(p2, alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r2 = p2.add_run("基于知识蒸馏的轻量化多模态感知模型")
    set_font(r2, SIMSUN, TIMES_NEW_ROMAN, 15, bold=True)  # 小三 = 15pt

    p3 = doc.add_paragraph()
    set_para_format(p3, alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r3 = p3.add_run("陶可  计算机z2204班  指导教师：李爱萍")
    set_font(r3, SIMSUN, TIMES_NEW_ROMAN, 12)

    # page break before 摘要
    para_pb = doc.add_paragraph()
    run_pb = para_pb.add_run()
    run_pb._r.append(docx_break_type())

    # ---- 摘要 ----
    add_heading_unnumbered(doc, "摘要", size=15)

    abstract_p1 = (
        "随着深度学习技术的迅猛发展，多模态感知模型在自动驾驶、智能交通等领域展现出显著性能优势，"
        "但其参数规模庞大、计算复杂度高，难以满足边缘设备实时部署的需求。知识蒸馏作为模型压缩领域的核心技术，"
        "通过将大规模教师模型所积累的知识迁移至轻量级学生模型，能够在大幅降低模型参数量与计算开销的同时有效保留"
        "原模型的性能表现。本文以基于知识蒸馏的轻量化多模态感知模型为研究背景，聚焦输出层知识蒸馏在"
        "Query-based多模态检测模型上的适配性问题。"
    )
    abstract_p2 = (
        "本文系统综述了多模态特征融合技术、知识蒸馏方法体系与多模态模型轻量化三个方向的研究现状。"
        "在特征融合方面，梳理了从两阶段CNN检测器到DETR系列端到端检测方法的演进，以及FUTR3D、DPFT等主流"
        "摄像头-雷达融合框架；在知识蒸馏方面，覆盖了经典软标签蒸馏、教师助手策略、解耦蒸馏（DKD）、"
        "关系蒸馏（RKD）与中间特征蒸馏等代表性方法；在多模态轻量化应用方面，分析了RayD3D、MultiDistiller"
        "等典型工作，并指出了现有方法在输出层蒸馏设计、Query对齐策略和雷达模态适配方面的不足。"
        "本综述为以DPFT模型为基础的知识蒸馏研究提供了系统性的理论支撑。"
    )
    add_body_para(doc, abstract_p1, indent_first=True)
    add_body_para(doc, abstract_p2, indent_first=True)

    kw_para = doc.add_paragraph()
    set_para_format(kw_para, indent_first=False)
    r_kw = kw_para.add_run("关键词：知识蒸馏；多模态感知；目标检测；轻量化；特征融合")
    set_font(r_kw, SIMSUN, TIMES_NEW_ROMAN, 12, bold=False)

    # page break before 目录
    para_pb2 = doc.add_paragraph()
    run_pb2 = para_pb2.add_run()
    run_pb2._r.append(docx_break_type())

    # ---- 目录 ----
    add_heading_unnumbered(doc, "目录", size=15)

    toc_entries = [
        ("摘要", ""),
        ("1  引言", ""),
        ("2  多模态感知模型与特征融合技术", ""),
        ("    2.1  目标检测基础方法", ""),
        ("    2.2  多模态特征融合方法", ""),
        ("    2.3  自动驾驶感知数据集", ""),
        ("    2.4  本章小结", ""),
        ("3  知识蒸馏技术", ""),
        ("    3.1  经典输出层知识蒸馏", ""),
        ("    3.2  改进的知识蒸馏方法", ""),
        ("    3.3  基于中间特征的知识蒸馏", ""),
        ("    3.4  本章小结", ""),
        ("4  多模态模型轻量化研究", ""),
        ("    4.1  多模态三维目标检测的知识蒸馏", ""),
        ("    4.2  现有方法的局限性与挑战", ""),
        ("    4.3  本章小结", ""),
        ("5  总结", ""),
        ("    5.1  本文主要工作", ""),
        ("    5.2  研究展望", ""),
        ("参考文献", ""),
    ]

    # Insert Word TOC field
    add_toc_field(doc)

    # Also add manual plain-text TOC as fallback
    for entry, _ in toc_entries:
        p_toc = doc.add_paragraph()
        set_para_format(p_toc, line_spacing=1.5, indent_first=False,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT)
        r_toc = p_toc.add_run(entry + " " + "." * max(1, 50 - len(entry)) + " x")
        set_font(r_toc, SIMSUN, TIMES_NEW_ROMAN, 12)

    # page break before Chapter 1
    para_pb3 = doc.add_paragraph()
    run_pb3 = para_pb3.add_run()
    run_pb3._r.append(docx_break_type())

    # ---- Chapter 1 ----
    add_heading1(doc, "1  引言")
    add_body_para(doc,
        "本章阐述本文的研究背景与研究问题，从多模态感知模型面临的部署挑战出发，引出知识蒸馏作为轻量化"
        "手段的研究价值，并概述本文的综述结构与主要内容安排。",
        indent_first=True)
    for t in ch1_paras:
        add_body_para(doc, t, indent_first=True)

    # ---- Chapter 2 ----
    para_pb4 = doc.add_paragraph()
    run_pb4 = para_pb4.add_run()
    run_pb4._r.append(docx_break_type())

    add_heading1(doc, "2  多模态感知模型与特征融合技术")
    add_body_para(doc,
        "本章围绕多模态感知模型的基础技术展开综述，依次介绍目标检测基础方法的发展脉络、多模态特征融合的"
        "主要技术范式，以及支撑模型研究的主流自动驾驶感知数据集，为后续章节的知识蒸馏与轻量化分析奠定背景基础。",
        indent_first=True)

    add_heading2(doc, "2.1  目标检测基础方法")
    add_body_para(doc, sec21, indent_first=True)

    add_heading2(doc, "2.2  多模态特征融合方法")
    for t in sec22_paras:
        add_body_para(doc, t, indent_first=True)

    add_heading2(doc, "2.3  自动驾驶感知数据集")
    add_body_para(doc, sec23, indent_first=True)

    add_heading2(doc, "2.4  本章小结")
    add_body_para(doc,
        "本章系统梳理了多模态感知模型与特征融合技术的研究现状。首先介绍了目标检测的基础方法，"
        "回顾了从Faster R-CNN、DeepLab到DETR的技术演进；其次综述了早期融合、特征层融合与决策层融合"
        "三种多模态特征融合范式，重点分析了FUTR3D和DPFT等代表性融合模型；最后梳理了KITTI、nuScenes、"
        "K-Radar等主流自动驾驶多模态数据集的特性与应用范围。上述技术构成了本文后续研究的基础框架。",
        indent_first=True)

    # ---- Chapter 3 ----
    para_pb5 = doc.add_paragraph()
    run_pb5 = para_pb5.add_run()
    run_pb5._r.append(docx_break_type())

    add_heading1(doc, "3  知识蒸馏技术")
    add_body_para(doc,
        "本章系统梳理知识蒸馏技术的研究进展，从经典输出层蒸馏方法出发，逐步深入改进型蒸馏策略与"
        "基于中间特征的蒸馏方法，重点关注各类方法的核心思想、技术特点及其在不同任务场景下的适用性分析。",
        indent_first=True)

    add_heading2(doc, "3.1  经典输出层知识蒸馏")
    for t in sec31:
        add_body_para(doc, t, indent_first=True)

    add_heading2(doc, "3.2  改进的知识蒸馏方法")
    for t in sec32:
        add_body_para(doc, t, indent_first=True)

    add_heading2(doc, "3.3  基于中间特征的知识蒸馏")
    for t in sec33:
        add_body_para(doc, t, indent_first=True)

    add_heading2(doc, "3.4  本章小结")
    add_body_para(doc,
        "本章对知识蒸馏技术体系进行了系统综述。从Hinton等提出的经典软标签蒸馏出发，介绍了基于响应、"
        "特征和关系三类知识的蒸馏方法；重点分析了教师助手策略、解耦知识蒸馏（DKD）、关系知识蒸馏（RKD）"
        "等改进方法；并进一步探讨了FitNets、特征激活边界蒸馏等基于中间特征的蒸馏方法。综合来看，现有蒸馏"
        "方法在分类任务上已较为成熟，但在Query-based多模态检测场景中的适配性研究仍存在明显空白，"
        "这为本课题的研究奠定了必要性基础。",
        indent_first=True)

    # ---- Chapter 4 ----
    para_pb6 = doc.add_paragraph()
    run_pb6 = para_pb6.add_run()
    run_pb6._r.append(docx_break_type())

    add_heading1(doc, "4  多模态模型轻量化研究")
    add_body_para(doc,
        "本章聚焦知识蒸馏技术在多模态三维目标检测场景中的具体应用，综述近年来代表性研究工作，"
        "并从方法设计、模态适配与工程实现多个维度分析现有方法的局限性，为本课题的研究方向提供明确的问题导向。",
        indent_first=True)

    add_heading2(doc, "4.1  多模态三维目标检测的知识蒸馏")
    for t in sec41:
        add_body_para(doc, t, indent_first=True)

    add_heading2(doc, "4.2  现有方法的局限性与挑战")
    add_body_para(doc, sec42, indent_first=True)

    add_heading2(doc, "4.3  本章小结")
    add_body_para(doc,
        "本章聚焦多模态感知模型的轻量化研究，梳理了知识蒸馏在多模态三维目标检测中的应用进展，"
        "包括RayD3D的射线深度蒸馏方法和MultiDistiller的多阶段渐进蒸馏框架，并综合分析了现有方法在"
        "输出层蒸馏设计、Query对齐策略和雷达模态适配三个维度上的局限性。上述挑战为本课题设计面向"
        "DPFT模型的输出层蒸馏方案提供了直接的问题导向。",
        indent_first=True)

    # ---- Chapter 5 ----
    para_pb7 = doc.add_paragraph()
    run_pb7 = para_pb7.add_run()
    run_pb7._r.append(docx_break_type())

    add_heading1(doc, "5  总结")
    add_heading2(doc, "5.1  本文主要工作")
    add_body_para(doc, ch5_paras[0], indent_first=True)
    add_body_para(doc, ch5_paras[1], indent_first=True)

    add_heading2(doc, "5.2  研究展望")
    add_body_para(doc,
        "基于上述综述分析，未来研究可在以下方向进一步深入：针对多模态输出层蒸馏的损失函数设计，"
        "探索结合匈牙利匹配的Query级别对齐策略；研究面向雷达模态稀疏特性的自适应蒸馏权重机制；"
        "以及探索渐进式蒸馏训练策略以缓解师生模型能力差距。",
        indent_first=True)

    # ---- 参考文献 ----
    para_pb8 = doc.add_paragraph()
    run_pb8 = para_pb8.add_run()
    run_pb8._r.append(docx_break_type())

    add_heading_unnumbered(doc, "参考文献", size=15)

    for ref in refs:
        p_ref = doc.add_paragraph()
        set_para_format(p_ref, line_spacing=1.5, indent_first=False,
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        # Prevent word break across lines
        pPr = p_ref._p.get_or_add_pPr()
        wrd = OxmlElement("w:wordWrap")
        wrd.set(qn("w:val"), "0")
        pPr.append(wrd)
        r_ref = p_ref.add_run(ref)
        set_font(r_ref, SIMSUN, TIMES_NEW_ROMAN, 12)

    out_path = "/home/runner/work/dpft/dpft/文献综述.docx"
    doc.save(out_path)
    print(f"✓ Saved {out_path}")


# ============================================================
# Build 外文文献及翻译.docx
# ============================================================
def build_waiwenxian():
    src = Document("/tmp/外文文献及翻译.docx")
    paras = [p.text.strip() for p in src.paragraphs]

    # Extract content sections
    bib_paras = [paras[3], paras[4], paras[5], paras[6]]
    en_note = paras[8]
    en_text = paras[9]
    cn_text = paras[11]
    glossary_items = [paras[i] for i in range(13, 18)]
    summary_text = paras[19]

    doc = Document()
    set_page_margins(doc)

    # ---- Cover ----
    p = doc.add_paragraph()
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r = p.add_run("外文文献及翻译")
    set_font(r, SIMHEI, TIMES_NEW_ROMAN, 22, bold=True)

    p2 = doc.add_paragraph()
    set_para_format(p2, alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    r2 = p2.add_run("学生：陶可    班级：计算机z2204班    日期：2026-04-15")
    set_font(r2, SIMSUN, TIMES_NEW_ROMAN, 12)

    # page break before Abstract
    pb = doc.add_paragraph()
    pb.add_run()._r.append(docx_break_type())

    # ---- Abstract (English) ----
    add_heading_unnumbered(doc, "Abstract", size=15)

    abs_p1 = (
        "This paper reviews the foreign literature titled "
        "\"DPFT: Dual Perspective Fusion Transformer for Camera-Radar-based Object Detection\" "
        "published in IEEE Transactions on Intelligent Vehicles (2025). "
        "The DPFT model addresses the challenge of robust 3D object detection in autonomous driving "
        "by proposing a dual-perspective fusion design that jointly optimizes bird's-eye-view and "
        "front-view representations using camera and 4D radar data."
    )
    abs_p2 = (
        "This document provides a selection of the original English text, a complete Chinese translation, "
        "a glossary of key technical terms, and a reading summary. "
        "The selected content focuses on the core fusion mechanism of DPFT, which iteratively exchanges "
        "complementary information between two perspectives to improve detection robustness in adverse "
        "weather and long-range scenarios. "
        "This foreign literature review is directly relevant to the research topic of lightweight "
        "multimodal perception via knowledge distillation, as DPFT serves as the teacher model in "
        "subsequent distillation experiments."
    )
    add_body_para(doc, abs_p1, font_name=TIMES_NEW_ROMAN, en_font=TIMES_NEW_ROMAN,
                  indent_first=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_body_para(doc, abs_p2, font_name=TIMES_NEW_ROMAN, en_font=TIMES_NEW_ROMAN,
                  indent_first=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # page break before TOC
    pb2 = doc.add_paragraph()
    pb2.add_run()._r.append(docx_break_type())

    # ---- 目录 ----
    add_heading_unnumbered(doc, "目录", size=15)
    add_toc_field(doc)

    toc_entries = [
        "Abstract",
        "1  外文文献信息",
        "2  英文原文（节选）",
        "3  中文翻译",
        "4  术语对照",
        "5  阅读总结",
        "    5.1  小结",
    ]
    for entry in toc_entries:
        p_toc = doc.add_paragraph()
        set_para_format(p_toc, line_spacing=1.5, indent_first=False,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT)
        r_toc = p_toc.add_run(entry + " " + "." * max(1, 50 - len(entry)) + " x")
        set_font(r_toc, SIMSUN, TIMES_NEW_ROMAN, 12)

    # ---- Section 1 ----
    pb3 = doc.add_paragraph()
    pb3.add_run()._r.append(docx_break_type())

    add_heading1(doc, "1  外文文献信息")
    for t in bib_paras:
        add_body_para(doc, t, indent_first=False)

    # ---- Section 2 ----
    pb4 = doc.add_paragraph()
    pb4.add_run()._r.append(docx_break_type())

    add_heading1(doc, "2  英文原文（节选）")
    add_body_para(doc, en_note, font_name=SIMSUN, indent_first=False)
    add_body_para(doc, en_text, font_name=TIMES_NEW_ROMAN, en_font=TIMES_NEW_ROMAN,
                  indent_first=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---- Section 3 ----
    pb5 = doc.add_paragraph()
    pb5.add_run()._r.append(docx_break_type())

    add_heading1(doc, "3  中文翻译")
    add_body_para(doc, cn_text, indent_first=True)

    # ---- Section 4 ----
    pb6 = doc.add_paragraph()
    pb6.add_run()._r.append(docx_break_type())

    add_heading1(doc, "4  术语对照")
    for item in glossary_items:
        add_body_para(doc, item, indent_first=False)

    # ---- Section 5 ----
    pb7 = doc.add_paragraph()
    pb7.add_run()._r.append(docx_break_type())

    add_heading1(doc, "5  阅读总结")
    add_body_para(doc, summary_text, indent_first=True)

    add_heading2(doc, "5.1  小结")
    add_body_para(doc,
        "本次外文文献阅读围绕DPFT模型的双视角融合机制展开，系统理解了其在摄像头-雷达融合检测任务中的核心贡献。"
        "通过精读原文、翻译对照与术语整理，深化了对多模态特征融合与Transformer检测架构的理解，"
        "为后续以DPFT为基础开展知识蒸馏研究奠定了坚实的文献基础。",
        indent_first=True)

    out_path = "/home/runner/work/dpft/dpft/外文文献及翻译.docx"
    doc.save(out_path)
    print(f"✓ Saved {out_path}")


if __name__ == "__main__":
    build_wenxian_zongsu()
    build_waiwenxian()
    print("✓ Both documents reformatted successfully.")
